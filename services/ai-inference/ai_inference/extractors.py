from __future__ import annotations

import io
import logging
from dataclasses import dataclass


@dataclass
class ExtractionResult:
    text: str
    language: str | None
    source: str


PDF_MIME = "application/pdf"
DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
TEXT_EXACT_MIME = {"application/json", "application/csv"}
TEXT_PREFIX = "text/"
IMAGE_PREFIX = "image/"


def _looks_like_pdf(file_name: str, mime_type: str) -> bool:
    return mime_type.lower() == PDF_MIME or file_name.lower().endswith(".pdf")


def _looks_like_docx(file_name: str, mime_type: str) -> bool:
    return mime_type.lower() == DOCX_MIME or file_name.lower().endswith(".docx")


def _looks_like_text(mime_type: str) -> bool:
    normalized = mime_type.lower()
    return normalized.startswith(TEXT_PREFIX) or normalized in TEXT_EXACT_MIME


def _looks_like_image(mime_type: str) -> bool:
    return mime_type.lower().startswith(IMAGE_PREFIX)


def _decode_text_bytes(payload: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="ignore")


def extract_text_plain(payload: bytes) -> ExtractionResult:
    text = _decode_text_bytes(payload).strip()
    if not text:
        raise ValueError("empty_text")
    return ExtractionResult(text=text, language=None, source="plain_text")


def extract_text_pdf(payload: bytes) -> ExtractionResult:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("missing_pdf_dependency:pypdf") from exc

    reader = PdfReader(io.BytesIO(payload))
    pages: list[str] = []
    for page in reader.pages:
        content = page.extract_text() or ""
        if content.strip():
            pages.append(content.strip())

    text = "\n\n".join(pages).strip()
    if text:
        return ExtractionResult(text=text, language=None, source="pdf_text")

    ocr_text = _extract_text_pdf_ocr_fallback(payload).strip()
    if not ocr_text:
        raise ValueError("empty_text")
    return ExtractionResult(text=ocr_text, language="en", source="pdf_ocr")


def _extract_text_pdf_ocr_fallback(payload: bytes) -> str:
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except Exception as exc:
        raise RuntimeError("missing_pdf_ocr_dependency:pdf2image+pytesseract") from exc

    try:
        images = convert_from_bytes(payload)
    except Exception as exc:
        error_name = exc.__class__.__name__.lower()
        if "pdfinfo" in error_name or "poppler" in str(exc).lower():
            raise RuntimeError("missing_system_dependency:poppler") from exc
        raise RuntimeError("pdf_render_failed") from exc

    lines: list[str] = []
    for image in images:
        try:
            page_text = pytesseract.image_to_string(image).strip()
        except pytesseract.pytesseract.TesseractNotFoundError as exc:
            raise RuntimeError("missing_system_dependency:tesseract") from exc
        if page_text:
            lines.append(page_text)

    return "\n\n".join(lines)


def extract_text_docx(payload: bytes) -> ExtractionResult:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("missing_docx_dependency:python-docx") from exc

    document = Document(io.BytesIO(payload))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text.strip():
                lines.append(row_text.strip())

    text = "\n".join(lines).strip()
    if not text:
        raise ValueError("empty_text")
    return ExtractionResult(text=text, language=None, source="docx_text")


def extract_text_image(payload: bytes) -> ExtractionResult:
    try:
        from PIL import Image
        import pytesseract
    except Exception as exc:
        raise RuntimeError("missing_ocr_dependency:pytesseract+pillow") from exc

    image = Image.open(io.BytesIO(payload))
    try:
        text = pytesseract.image_to_string(image).strip()
    except pytesseract.pytesseract.TesseractNotFoundError as exc:
        raise RuntimeError("missing_system_dependency:tesseract") from exc

    # Many photos naturally contain no readable text. This should not be treated
    # as an extraction failure.
    return ExtractionResult(text=text, language="en", source="ocr_image")


def extract_text_for_file(
    file_name: str,
    mime_type: str,
    payload: bytes,
    logger: logging.Logger,
) -> ExtractionResult:
    if _looks_like_pdf(file_name, mime_type):
        logger.info("Extraction route: pdf file=%s", file_name)
        return extract_text_pdf(payload)

    if _looks_like_docx(file_name, mime_type):
        logger.info("Extraction route: docx file=%s", file_name)
        return extract_text_docx(payload)

    if _looks_like_text(mime_type):
        logger.info("Extraction route: plain-text file=%s mime=%s", file_name, mime_type)
        return extract_text_plain(payload)

    if _looks_like_image(mime_type):
        logger.info("Extraction route: image-ocr file=%s mime=%s", file_name, mime_type)
        return extract_text_image(payload)

    raise ValueError("unsupported_file_type")
